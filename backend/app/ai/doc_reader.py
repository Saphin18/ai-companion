"""
Document -> plain text (PDF, DOCX, TXT/MD).
No AI involved; this is pure parsing. Best-effort: returns None on failure.
"""
import io

MAX_BYTES = 5 * 1024 * 1024      # reject anything bigger before parsing
MAX_CHARS = 15000                # protects Render memory and Groq rate limits


def _truncate(text: str) -> str | None:
    text = (text or "").strip()
    if not text:
        return None
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[...document truncated...]"
    return text


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
        if sum(len(p) for p in parts) > MAX_CHARS:
            break
    return "\n".join(parts)


def _read_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(
    data: bytes,
    filename: str = "",
    mime_type: str = "",
) -> str | None:
    """Pull readable text out of a document. Returns None if it can't."""
    if not data or len(data) > MAX_BYTES:
        return None

    name = (filename or "").lower()
    mime = (mime_type or "").lower()

    try:
        if name.endswith(".pdf") or "pdf" in mime:
            return _truncate(_read_pdf(data))
        if name.endswith(".docx") or "wordprocessingml" in mime:
            return _truncate(_read_docx(data))
        if name.endswith((".txt", ".md", ".csv")) or mime.startswith("text/"):
            return _truncate(data.decode("utf-8", errors="replace"))
    except Exception:
        return None

    return None
